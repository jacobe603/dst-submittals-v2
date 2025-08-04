"""Tag extraction pipeline stage with switchable parsing modes"""

import os
import re
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

from pipeline.base import PipelineStage, PipelineContext, StageResult


class TagExtractionStage(PipelineStage):
    """
    Extract equipment tags from documents using configurable parsing modes.
    
    Input: List of file paths in context['files']
    Output: Dict mapping files to extracted tags with confidence scores
    Config: 
        - mode: 'content' or 'filename' 
        - confidence_threshold: Minimum confidence for tag acceptance
        - enable_filename_fallback: Use filename parsing if content fails
    """
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__("tag_extraction", config)
        
        # Extraction mode
        self.mode = config.get('mode', 'content')
        self.confidence_threshold = config.get('confidence_threshold', 0.8)
        self.enable_filename_fallback = config.get('enable_filename_fallback', True)
        
        # Tag patterns for content extraction
        self.content_patterns = [
            # Direct tag patterns
            (r'Unit Tag:\s*([A-Z]{2,4}-[0-9]+[A-Z]*)', 0.95),
            (r'Tag:\s*([A-Z]{2,4}-[0-9]+[A-Z]*)', 0.9),
            (r'Unit:\s*([A-Z]{2,4}-[0-9]+[A-Z]*)', 0.85),
            
            # Equipment patterns
            (r'([A-Z]{2,4}-[0-9]+[A-Z]*)\s+(?:Unit|Equipment)', 0.8),
            (r'Equipment\s+([A-Z]{2,4}-[0-9]+[A-Z]*)', 0.8),
            
            # Table/form patterns
            (r'([A-Z]{2,4}-[0-9]+[A-Z]*)\s*[:\-]\s*(?:Air Handler|MAU|AHU)', 0.75),
            
            # Generic patterns (lower confidence)
            (r'\b([A-Z]{2,4}-[0-9]+[A-Z]*)\b', 0.6),
        ]
        
        # Filename patterns
        self.filename_patterns = [
            # Standard format: "10_Item Summary.docx" -> AHU-10
            (r'^(\d+)_.*', self._number_to_ahu, 0.8),
            
            # Direct tag format: "AHU-10_Technical_Data.docx"
            (r'^([A-Z]{2,4}-\d+[A-Z]*)_.*', lambda m: m.group(1), 0.9),
            
            # Tag with spaces: "AHU 10 Drawing.docx"
            (r'^([A-Z]{2,4})\s*(\d+[A-Z]*).*', lambda m: f"{m.group(1)}-{m.group(2)}", 0.85),
        ]
        
    def validate_input(self, context: PipelineContext) -> bool:
        """Validate that we have files to process"""
        files = context.get('files', [])
        if not files:
            return False
            
        # Check that files exist
        for file_path in files:
            if not os.path.exists(file_path):
                if self.logger:
                    self.logger.error(f"File not found: {file_path}")
                return False
                
        return True
        
    def process(self, context: PipelineContext) -> StageResult:
        """Extract tags from all files using configured mode"""
        files = context.get('files', [])
        tag_mapping = {}
        extraction_details = {}
        
        if self.logger:
            self.logger.info(f"Extracting tags from {len(files)} files using {self.mode} mode")
            
        for file_path in files:
            try:
                filename = os.path.basename(file_path)
                
                if self.logger:
                    self.logger.debug(f"Processing file: {filename}")
                
                # Extract tag based on mode
                if self.mode == 'filename':
                    tag, confidence, details = self._extract_from_filename(filename)
                    
                    # Fallback to content if enabled and filename extraction failed
                    if not tag and self.enable_filename_fallback:
                        tag_content, conf_content, details_content = self._extract_from_content(file_path)
                        if tag_content and conf_content > confidence:
                            tag, confidence, details = tag_content, conf_content, details_content
                            details['fallback_used'] = 'content'
                            
                elif self.mode == 'content':
                    tag, confidence, details = self._extract_from_content(file_path)
                    
                    # Fallback to filename if enabled and content extraction failed
                    if not tag and self.enable_filename_fallback:
                        tag_filename, conf_filename, details_filename = self._extract_from_filename(filename)
                        if tag_filename and conf_filename > confidence:
                            tag, confidence, details = tag_filename, conf_filename, details_filename
                            details['fallback_used'] = 'filename'
                            
                else:
                    raise ValueError(f"Unknown extraction mode: {self.mode}")
                
                # Store results
                extraction_details[filename] = {
                    'tag': tag,
                    'confidence': confidence,
                    'details': details,
                    'accepted': confidence >= self.confidence_threshold
                }
                
                if tag and confidence >= self.confidence_threshold:
                    tag_mapping[filename] = tag
                    if self.logger:
                        self.logger.info(f"Extracted tag {tag} from {filename} (confidence: {confidence:.2f})")
                else:
                    if self.logger:
                        if tag:
                            self.logger.warning(
                                f"Tag {tag} rejected for {filename} (confidence: {confidence:.2f} < {self.confidence_threshold})"
                            )
                        else:
                            self.logger.warning(f"No tag found in {filename}")
                            
            except Exception as e:
                if self.logger:
                    self.logger.error(f"Error processing {filename}: {e}")
                extraction_details[filename] = {
                    'tag': None,
                    'confidence': 0.0,
                    'error': str(e),
                    'accepted': False
                }
                
        # Log summary
        if self.logger:
            accepted_count = sum(1 for d in extraction_details.values() if d['accepted'])
            self.logger.info(f"Tag extraction complete: {accepted_count}/{len(files)} files processed successfully")
            
        return StageResult(
            success=True,
            data={
                'tag_mapping': tag_mapping,
                'extraction_details': extraction_details
            },
            debug_info={
                'mode': self.mode,
                'confidence_threshold': self.confidence_threshold,
                'total_files': len(files),
                'successful_extractions': len(tag_mapping)
            }
        )
        
    def _extract_from_filename(self, filename: str) -> Tuple[Optional[str], float, Dict[str, Any]]:
        """Extract tag from filename using patterns"""
        details = {'method': 'filename', 'patterns_tried': []}
        
        for pattern, transform_func, base_confidence in self.filename_patterns:
            match = re.match(pattern, filename, re.IGNORECASE)
            details['patterns_tried'].append({
                'pattern': pattern,
                'matched': bool(match)
            })
            
            if match:
                try:
                    tag = transform_func(match)
                    tag = self._normalize_tag(tag)
                    
                    details['matched_pattern'] = pattern
                    details['raw_match'] = match.groups()
                    details['normalized_tag'] = tag
                    
                    return tag, base_confidence, details
                    
                except Exception as e:
                    details['transform_error'] = str(e)
                    continue
                    
        return None, 0.0, details
        
    def _extract_from_content(self, file_path: str) -> Tuple[Optional[str], float, Dict[str, Any]]:
        """Extract tag from file content using patterns"""  
        details = {'method': 'content', 'patterns_tried': []}
        
        try:
            # Read file content
            content = self._read_file_content(file_path)
            if not content:
                details['error'] = 'Could not read file content'
                return None, 0.0, details
                
            details['content_length'] = len(content)
            
            # Try patterns in order of confidence
            best_tag = None
            best_confidence = 0.0
            
            for pattern, base_confidence in self.content_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                pattern_info = {
                    'pattern': pattern,
                    'matches': matches,
                    'confidence': base_confidence
                }
                details['patterns_tried'].append(pattern_info)
                
                if matches:
                    # Use first match for now (could be improved with frequency analysis)
                    raw_tag = matches[0]
                    normalized_tag = self._normalize_tag(raw_tag)
                    
                    if base_confidence > best_confidence:
                        best_tag = normalized_tag
                        best_confidence = base_confidence
                        details['best_match'] = {
                            'pattern': pattern,
                            'raw_tag': raw_tag,
                            'normalized_tag': normalized_tag
                        }
                        
            return best_tag, best_confidence, details
            
        except Exception as e:
            details['error'] = str(e)
            return None, 0.0, details
            
    def _read_file_content(self, file_path: str) -> Optional[str]:
        """Read content from various file types"""
        try:
            if file_path.endswith('.docx'):
                return self._read_docx_content(file_path)
            elif file_path.endswith('.doc'):
                return self._read_doc_content(file_path)
            elif file_path.endswith('.pdf'):
                return self._read_pdf_content(file_path)
            else:
                # Try reading as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            if self.logger:
                self.logger.debug(f"Could not read content from {file_path}: {e}")
            return None
            
    def _read_docx_content(self, file_path: str) -> Optional[str]:
        """Read content from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
                
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content.append(cell.text)
                        
            return '\n'.join(content)
            
        except ImportError:
            if self.logger:
                self.logger.debug("python-docx not available for DOCX reading")
            return None
        except Exception as e:
            if self.logger:
                self.logger.debug(f"Error reading DOCX {file_path}: {e}")
            return None
            
    def _read_doc_content(self, file_path: str) -> Optional[str]:
        """Read content from DOC file using various methods"""
        # This would use the same logic as the existing enhanced_doc_extractor
        # For now, return None to fallback to filename
        return None
        
    def _read_pdf_content(self, file_path: str) -> Optional[str]:
        """Read content from PDF file"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            content = []
            
            for page in reader.pages:
                content.append(page.extract_text())
                
            return '\n'.join(content)
            
        except ImportError:
            if self.logger:
                self.logger.debug("pypdf not available for PDF reading")
            return None
        except Exception as e:
            if self.logger:
                self.logger.debug(f"Error reading PDF {file_path}: {e}")
            return None
            
    def _number_to_ahu(self, match) -> str:
        """Convert number to AHU tag (e.g., 10 -> AHU-10)"""
        number = match.group(1)
        
        # Determine if it should be MAU or AHU based on context
        # For now, assume AHU (could be enhanced with more logic)
        if int(number) >= 12:  # Common convention: MAU-12+
            return f"MAU-{number}"
        else:
            return f"AHU-{number}"
            
    def _normalize_tag(self, tag: str) -> str:
        """Normalize tag format (e.g., AHU01 -> AHU-1)"""
        # Remove extra spaces
        tag = tag.strip().upper()
        
        # Ensure hyphen separation
        match = re.match(r'([A-Z]+)[-\s]*(\d+[A-Z]*)', tag)
        if match:
            prefix, number = match.groups()
            
            # Remove leading zeros from number part
            number = re.sub(r'^0+(\d)', r'\1', number)
            
            return f"{prefix}-{number}"
            
        return tag
        
    def validate_output(self, result: StageResult) -> bool:
        """Validate that we produced a valid tag mapping"""
        if not result.success:
            return True  # Let error be handled upstream
            
        tag_mapping = result.data.get('tag_mapping', {})
        
        # Check that tag mapping is a dict
        if not isinstance(tag_mapping, dict):
            return False
            
        # Check that all tags follow expected format
        tag_pattern = re.compile(r'^[A-Z]{2,4}-\d+[A-Z]*$')
        for filename, tag in tag_mapping.items():
            if not tag_pattern.match(tag):
                if self.logger:
                    self.logger.warning(f"Invalid tag format: {tag} for file {filename}")
                    
        return True