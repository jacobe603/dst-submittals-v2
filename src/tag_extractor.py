#!/usr/bin/env python3
"""
Comprehensive tag extraction script for CS Air Handler documents
Handles both .doc and .docx files using multiple extraction methods
"""

import os
import re
import json
import glob
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
import docx2txt
from PIL import Image

# Import diagnostic logging functions
try:
    from .logger import log_tag_extraction, log_json_snapshot, log_processing_stage
except ImportError:
    from logger import log_tag_extraction, log_json_snapshot, log_processing_stage

class TagExtractor:
    def __init__(self, docs_path: str, use_filename_tags: bool = False):
        """
        Initialize TagExtractor with configurable extraction method.
        
        Args:
            docs_path: Path to directory containing documents
            use_filename_tags: If True, extract tags from filenames only (fast).
                             If False, use content-based extraction (slower but more flexible).
        """
        self.docs_path = docs_path
        self.tag_mapping = {}
        self.extraction_log = []
        self.use_filename_tags = use_filename_tags
        
    def log_extraction(self, filename: str, method: str, success: bool, tag: str = None, error: str = None):
        """Log extraction attempts"""
        log_entry = {
            'filename': filename,
            'method': method,
            'success': success,
            'tag': tag,
            'error': error
        }
        self.extraction_log.append(log_entry)
        
    def extract_tag_from_filename(self, filename: str) -> Optional[str]:
        """
        Extract tag directly from filename by finding text before document type suffixes.
        
        This method provides fast, reliable tag extraction without reading file content.
        It extracts the raw text before common document type patterns, allowing users
        to manually edit the extracted tags in the structure editor for perfect accuracy.
        
        Performance benefits:
        - No file I/O required (instant extraction)
        - Works with any file type/format 
        - Reliable even with corrupted/locked files
        
        Examples:
            "10_Technical_Data_Sheet.docx" → "10"
            "AHU-E1_Drawing.doc" → "AHU-E1"  
            "AHU-10 - Technical Data Sheet.docx" → "AHU-10"
            "AHU-M3 3-20-2025 - Fan Curve.doc" → "AHU-M3 3-20-2025"  
            "28_Fan_Curve_Supply.jpg" → "28_Fan_Curve"
            "MAU-12_Item_Summary.docx" → "MAU-12"
            
        Args:
            filename: The filename to extract tag from
            
        Returns:
            Raw extracted text (no formatting applied) or None if no pattern found
        """
        # Remove file extension first
        base_name = os.path.splitext(filename)[0]
        
        # Handle new format with " - " separator (e.g., "AHU-10 - Technical Data Sheet")
        # This is the primary format for tagged files
        dash_separated_patterns = [
            ' - PreciseLine Drawings',
            ' - Technical Data Sheet',
            ' - Item Summary',
            ' - Drawing', 
            ' - Fan Curve',
            ' - Supply',  # For fan curves
            ' - Return',
            ' - Exhaust'
        ]
        
        for pattern in dash_separated_patterns:
            if pattern in base_name:
                # Extract everything before the pattern
                tag = base_name.split(pattern)[0]
                if tag.strip():
                    return tag.strip()
        
        # Handle legacy format with underscore separators
        # Define common document type suffixes to extract text before
        # IMPORTANT: Order matters - longer patterns first to avoid partial matches
        underscore_suffixes = [
            '_PreciseLine_Drawings',  # Must come before _PreciseLine
            '_-_Technical_Data_Sheet', # Handle _-_ separator pattern
            '_-_Item_Summary',         # Handle _-_ separator pattern
            '_-_Drawing',              # Handle _-_ separator pattern
            '_-_Fan_Curve',            # Handle _-_ separator pattern
            '_Technical_Data_Sheet',
            '_Item_Summary', 
            '_Drawing',
            '_Fan_Curve',
            '_PreciseLine',
            '_UnitCADFile'
        ]
        
        # Try to find text before each suffix pattern
        for suffix in underscore_suffixes:
            if suffix in base_name:
                # Extract everything before the suffix
                tag = base_name.split(suffix)[0]
                if tag.strip():  # Make sure we have actual content
                    # Clean up common filename artifacts
                    # Remove trailing _- patterns (e.g., "AHU-E3_-" -> "AHU-E3")
                    tag = tag.rstrip('_-').strip()
                    return tag if tag else None
        
        # If no suffix pattern found, try common separators
        # This handles cases like "28_Fan_Curve_Supply" where we want "28_Fan_Curve"
        separators = ['_Supply', '_Return', '_Exhaust']
        for sep in separators:
            if base_name.endswith(sep):
                tag = base_name[:-len(sep)]
                if tag.strip():
                    return tag.strip()
        
        # Final fallback: if no patterns match, use the whole filename without extension
        # This ensures we always return something for manual editing
        return base_name if base_name.strip() else None
        
    def extract_tags_from_text(self, text: str, filename: str) -> List[str]:
        """Extract all possible tags from text content"""
        tags = []
        
        # Common tag patterns
        patterns = [
            r'Unit Tag:\s*([A-Z0-9\-\s]+)',  # Unit Tag: AHU-10
            r'Unit Tag\s+([A-Z0-9\-\s]+)',   # Unit Tag AHU-10
            r'(AHU-[A-Z0-9\-\s]+)',          # AHU-10, AHU-E1, etc.
            r'(MAU-[A-Z0-9\-\s]+)',          # MAU-12, etc.
            r'Unit:\s*([A-Z0-9\-\s]+)',      # Unit: AHU-10
            r'Tag:\s*([A-Z0-9\-\s]+)'        # Tag: AHU-10
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Clean up the match
                clean_tag = match.strip().upper()
                # Remove common suffixes and clean formatting
                clean_tag = re.sub(r'\s+', ' ', clean_tag)  # Multiple spaces to single
                clean_tag = clean_tag.split()[0] if clean_tag else ''  # Take first word
                
                if clean_tag and (clean_tag.startswith('AHU-') or clean_tag.startswith('MAU-')):
                    if clean_tag not in tags:
                        tags.append(clean_tag)
        
        return tags
    
    def convert_image_to_pdf(self, image_path: str, output_dir: str = None) -> Optional[str]:
        """Convert JPG/PNG image to PDF"""
        try:
            if output_dir is None:
                output_dir = os.path.dirname(image_path)
            
            # Create output filename
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
            
            # Convert image to PDF
            image = Image.open(image_path)
            
            # Convert to RGB if necessary (for PNG with transparency)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Save as PDF
            image.save(pdf_path, 'PDF', resolution=100.0, quality=95)
            
            print(f"  [CONVERTED] {os.path.basename(image_path)} -> {os.path.basename(pdf_path)}")
            return pdf_path
            
        except Exception as e:
            print(f"  [ERROR] Failed to convert {os.path.basename(image_path)}: {str(e)}")
            return None
    
    def extract_filename_prefix(self, filename: str) -> Optional[str]:
        """Extract the numeric prefix from filename (e.g., '10_' from '10_Fan Curve - Supply.jpg')"""
        import re
        match = re.match(r'^(\d+)_', filename)
        return match.group(1) if match else None
    
    def find_tag_by_filename_matching(self, target_filename: str) -> Optional[str]:
        """Find tag by matching filename prefix with other files in the directory"""
        target_prefix = self.extract_filename_prefix(target_filename)
        if not target_prefix:
            return None
            
        # Look for .docx files with the same prefix first (they have the best success rate)
        docx_files = glob.glob(os.path.join(self.docs_path, f"{target_prefix}_*.docx"))
        for docx_file in docx_files:
            tag = self.extract_from_docx_python_docx(docx_file)
            if tag:
                self.log_extraction(target_filename, 'filename_matching_docx', True, tag)
                return tag
        
        # Fallback to .doc files with the same prefix
        doc_files = glob.glob(os.path.join(self.docs_path, f"{target_prefix}_*.doc"))
        for doc_file in doc_files:
            # Try string extraction method
            tag = self.extract_from_doc_strings(doc_file)
            if tag:
                self.log_extraction(target_filename, 'filename_matching_doc', True, tag)
                return tag
                
        return None
    

    def extract_from_docx_python_docx(self, file_path: str) -> Optional[str]:
        """Extract tag using python-docx library"""
        try:
            doc = Document(file_path)
            
            # Extract all text from paragraphs
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += '\n' + cell.text
            
            tags = self.extract_tags_from_text(full_text, os.path.basename(file_path))
            return tags[0] if tags else None
            
        except Exception as e:
            self.log_extraction(os.path.basename(file_path), 'python-docx', False, error=str(e))
            return None
    
    def extract_from_docx_docx2txt(self, file_path: str) -> Optional[str]:
        """Extract tag using docx2txt library"""
        try:
            text = docx2txt.process(file_path)
            tags = self.extract_tags_from_text(text, os.path.basename(file_path))
            return tags[0] if tags else None
            
        except Exception as e:
            self.log_extraction(os.path.basename(file_path), 'docx2txt', False, error=str(e))
            return None
    
    def extract_from_doc_olefile(self, file_path: str) -> Optional[str]:
        """Attempt to extract tag from .doc file using olefile (limited success expected)"""
        try:
            import olefile
            
            if not olefile.isOleFile(file_path):
                return None
                
            ole = olefile.OleFileIO(file_path)
            
            # Try to find text streams
            text_content = ""
            
            # Look for common text streams in Word documents
            possible_streams = ['WordDocument', '1Table', '0Table', 'Data']
            
            for stream_name in possible_streams:
                if ole._olestream_size.get(stream_name):
                    try:
                        stream = ole.opendir(stream_name)
                        # This is a simplified approach - real .doc parsing is much more complex
                        raw_data = ole._olestream[stream_name]
                        # Try to extract readable text (very basic approach)
                        text_content += raw_data.decode('utf-8', errors='ignore')
                    except:
                        continue
            
            ole.close()
            
            if text_content:
                tags = self.extract_tags_from_text(text_content, os.path.basename(file_path))
                return tags[0] if tags else None
                
        except Exception as e:
            self.log_extraction(os.path.basename(file_path), 'olefile', False, error=str(e))
            return None
    
    def extract_from_doc_strings(self, file_path: str) -> Optional[str]:
        """Extract readable strings from .doc file (basic approach)"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Convert to string and extract readable parts
            text = content.decode('utf-8', errors='ignore')
            
            # Look for tag patterns in the raw text
            tags = self.extract_tags_from_text(text, os.path.basename(file_path))
            return tags[0] if tags else None
            
        except Exception as e:
            self.log_extraction(os.path.basename(file_path), 'strings', False, error=str(e))
            return None
    
    def extract_tag_from_file(self, file_path: str) -> Optional[str]:
        """
        Extract tag from a single file using either filename-based or content-based methods.
        
        When use_filename_tags is True:
        - Extracts tag directly from filename (fast, no file I/O)
        - Logs extraction method for transparency
        
        When use_filename_tags is False:
        - Uses traditional content-based extraction (slower but flexible)
        - Falls back through multiple extraction methods
        """
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()
        
        print(f"Processing: {filename}")
        
        # Use filename-based extraction if enabled
        if self.use_filename_tags:
            tag = self.extract_tag_from_filename(filename)
            if tag:
                self.log_extraction(filename, 'filename', True, tag)
                # Enhanced diagnostic logging
                log_tag_extraction(filename, 'filename', True, tag, ['filename_patterns'])
                print(f"  [FILENAME] {filename} -> {tag}")
                return tag
            else:
                self.log_extraction(filename, 'filename', False, error="No filename pattern matched")
                # Enhanced diagnostic logging
                log_tag_extraction(filename, 'filename', False, patterns_tested=['filename_patterns'])
                print(f"  [FILENAME] {filename} -> No pattern found")
                return None
        
        # Traditional content-based extraction
        methods_tested = []
        if file_ext == '.docx':
            # Try python-docx first
            methods_tested.append('python-docx')
            tag = self.extract_from_docx_python_docx(file_path)
            if tag:
                self.log_extraction(filename, 'python-docx', True, tag)
                log_tag_extraction(filename, 'python-docx', True, tag, methods_tested)
                return tag
            
            # Try docx2txt as fallback
            methods_tested.append('docx2txt')
            tag = self.extract_from_docx_docx2txt(file_path)
            if tag:
                self.log_extraction(filename, 'docx2txt', True, tag)
                log_tag_extraction(filename, 'docx2txt', True, tag, methods_tested)
                return tag
                
        elif file_ext == '.doc':
            # Try multiple approaches for .doc files
            
            # Method 1: Raw string extraction
            methods_tested.append('strings')
            tag = self.extract_from_doc_strings(file_path)
            if tag:
                self.log_extraction(filename, 'strings', True, tag)
                log_tag_extraction(filename, 'strings', True, tag, methods_tested)
                return tag
            
            # Method 2: olefile approach
            methods_tested.append('olefile')
            tag = self.extract_from_doc_olefile(file_path)
            if tag:
                self.log_extraction(filename, 'olefile', True, tag)
                log_tag_extraction(filename, 'olefile', True, tag, methods_tested)
                return tag
        
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            # Use filename matching for image files
            methods_tested.append('filename_matching')
            tag = self.find_tag_by_filename_matching(filename)
            if tag:
                log_tag_extraction(filename, 'filename_matching', True, tag, methods_tested)
                # Also convert the image to PDF
                self.convert_image_to_pdf(file_path)
                return tag
        
        # If no tag found, log the failure
        self.log_extraction(filename, 'all_methods', False, error='No tag found')
        log_tag_extraction(filename, 'all_methods', False, patterns_tested=methods_tested)
        return None
    
    def extract_all_tags(self) -> Dict[str, str]:
        """Extract tags from all files in the directory using configured method"""
        # Log the start of tag extraction
        log_processing_stage('extract_all_tags', 'started', {
            'docs_path': self.docs_path,
            'use_filename_tags': self.use_filename_tags
        })
        
        print("="*60)
        if self.use_filename_tags:
            print("EXTRACTING TAGS FROM FILENAMES (FAST MODE)")
            print("Note: Tags extracted from filenames can be edited in structure editor")
        else:
            print("EXTRACTING TAGS FROM FILE CONTENT (DEEP SCAN)")
            print("Note: Scanning file content for tag patterns")
        print("="*60)
        
        # Get all .doc, .docx, and image files
        doc_files = glob.glob(os.path.join(self.docs_path, "*.doc"))
        docx_files = glob.glob(os.path.join(self.docs_path, "*.docx"))
        image_files = glob.glob(os.path.join(self.docs_path, "*.jpg")) + \
                      glob.glob(os.path.join(self.docs_path, "*.jpeg")) + \
                      glob.glob(os.path.join(self.docs_path, "*.png"))
        
        all_files = sorted(doc_files + docx_files + image_files)
        
        # Log file discovery
        log_processing_stage('file_discovery', 'completed', {
            'docx_files': len(docx_files),
            'doc_files': len(doc_files),
            'image_files': len(image_files),
            'total_files': len(all_files)
        })
        
        print(f"Found {len(docx_files)} .docx files, {len(doc_files)} .doc files, and {len(image_files)} image files")
        print()
        
        # Process each file
        successful_extractions = 0
        failed_extractions = 0
        
        for file_path in all_files:
            filename = os.path.basename(file_path)
            tag = self.extract_tag_from_file(file_path)
            
            if tag:
                self.tag_mapping[filename] = tag
                successful_extractions += 1
                print(f"  [OK] {filename} -> {tag}")
            else:
                self.tag_mapping[filename] = None
                failed_extractions += 1
                print(f"  [FAIL] {filename} -> No tag found")
        
        # Log completion results
        log_processing_stage('extract_all_tags', 'completed', {
            'total_files': len(all_files),
            'successful_extractions': successful_extractions,
            'failed_extractions': failed_extractions,
            'success_rate': (successful_extractions / len(all_files) * 100) if all_files else 0
        })
        
        # Log JSON snapshot of tag mapping
        log_json_snapshot('tag_mapping_results', self.tag_mapping)
        
        return self.tag_mapping
    
    def create_tag_groups(self) -> Dict[str, List[str]]:
        """Group files by their tags"""
        tag_groups = {}
        
        for filename, tag in self.tag_mapping.items():
            if tag:
                if tag not in tag_groups:
                    tag_groups[tag] = []
                tag_groups[tag].append(filename)
        
        return tag_groups
    
    def save_results(self, output_file: str = "tag_mapping.json"):
        """Save tag mapping and results to JSON file"""
        results = {
            'tag_mapping': self.tag_mapping,
            'tag_groups': self.create_tag_groups(),
            'extraction_log': self.extraction_log,
            'summary': {
                'total_files': len(self.tag_mapping),
                'files_with_tags': len([f for f in self.tag_mapping.values() if f]),
                'files_without_tags': len([f for f in self.tag_mapping.values() if not f]),
                'unique_tags': len(set([t for t in self.tag_mapping.values() if t]))
            }
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        
        return output_file
    
    def print_summary(self):
        """Print extraction summary"""
        tag_groups = self.create_tag_groups()
        
        print("\n" + "="*60)
        print("TAG EXTRACTION SUMMARY")
        print("="*60)
        
        print(f"Total files processed: {len(self.tag_mapping)}")
        print(f"Files with tags found: {len([f for f in self.tag_mapping.values() if f])}")
        print(f"Files without tags: {len([f for f in self.tag_mapping.values() if not f])}")
        print(f"Unique tags found: {len(tag_groups)}")
        
        print("\nTags found (ordered):")
        for tag in sorted(tag_groups.keys()):
            files = tag_groups[tag]
            print(f"  {tag}: {len(files)} files")
            for filename in sorted(files):
                print(f"    - {filename}")
        
        print("\nFiles without tags:")
        for filename, tag in self.tag_mapping.items():
            if not tag:
                print(f"  - {filename}")

def main():
    docs_path = r"C:\Users\jacob\Claude\python-docx\documents\CS_Air_Handler_Light_Kit"
    
    extractor = TagExtractor(docs_path)
    extractor.extract_all_tags()
    
    # Save results
    output_file = extractor.save_results()
    print(f"\nResults saved to: {output_file}")
    
    # Print summary
    extractor.print_summary()
    
    return extractor

if __name__ == "__main__":
    extractor = main()