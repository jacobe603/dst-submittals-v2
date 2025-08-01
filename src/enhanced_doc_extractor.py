#!/usr/bin/env python3
"""
Enhanced document processor that generates complete PDF structure JSON
Creates a comprehensive structure including title pages, documents, and cut sheets
"""

import os
import re
import json
import glob
from typing import Optional, Dict, List
from pathlib import Path

# Import diagnostic logging functions
try:
    from .logger import log_processing_stage, log_json_snapshot
except ImportError:
    from logger import log_processing_stage, log_json_snapshot

def has_pricing_content(file_path: str) -> bool:
    """Check if a file contains pricing information ($ symbols)"""
    if not os.path.exists(file_path):
        return False
        
    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1].lower()
    
    # For image files, we can't easily check content, so check filename
    if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
        # Check if filename suggests pricing (e.g., "Item Summary", "Pricing", "Cost")
        pricing_keywords = ['item summary', 'pricing', 'cost', 'price', 'quote']
        filename_lower = filename.lower()
        return any(keyword in filename_lower for keyword in pricing_keywords)
    
    # For document files, try to read content
    try:
        if file_ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            # Check paragraphs
            for para in doc.paragraphs:
                if '$' in para.text:
                    return True
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '$' in cell.text:
                            return True
                            
        elif file_ext == '.doc':
            # For .doc files, do a simple binary search for $ symbols
            with open(file_path, 'rb') as f:
                content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            if '$' in text_content:
                return True
                
        elif file_ext == '.pdf':
            # For PDF files, we'll assume they don't have pricing unless filename suggests it
            pricing_keywords = ['item summary', 'pricing', 'cost', 'price', 'quote']
            filename_lower = filename.lower()
            return any(keyword in filename_lower for keyword in pricing_keywords)
            
    except Exception as e:
        print(f"  [WARNING] Could not check pricing content in {filename}: {e}")
        # If we can't read the file, check filename for pricing keywords
        pricing_keywords = ['item summary', 'pricing', 'cost', 'price', 'quote']
        filename_lower = filename.lower()
        return any(keyword in filename_lower for keyword in pricing_keywords)
    
    return False

def classify_file_type(filename: str) -> str:
    """Classify file type based on filename patterns"""
    filename_lower = filename.lower()
    
    # Normalize separators to make matching more robust
    # Replaces both underscores and hyphens with spaces
    normalized_name = filename_lower.replace('_', ' ').replace('-', ' ')
    
    # Remove numeric prefix for classification
    clean_name = re.sub(r'^\d+\s', '', normalized_name)

    # Perform checks on the clean, normalized name
    if 'technical data sheet' in clean_name or 'tech data' in clean_name:
        return 'Technical Data Sheet'
    elif 'fan curve' in clean_name:
        return 'Fan Curve'
    elif 'drawing' in clean_name:
        return 'Drawing'
    elif 'cut sheet' in clean_name or filename.upper().startswith('CS'):
        return 'Cut Sheet'
    elif 'item summary' in clean_name:
        return 'Item Summary'
    else:
        return 'Other'

def create_display_title(filename: str) -> str:
    """Create a clean display title from filename"""
    # Remove file extension
    clean_name = os.path.splitext(filename)[0]
    
    # Remove numeric prefix (e.g., "10_" from "10_Technical Data Sheet")
    clean_name = re.sub(r'^\d+_', '', clean_name)
    
    # Replace underscores and hyphens with spaces to be consistent with classify_file_type
    clean_name = clean_name.replace('_', ' ').replace('-', ' ')
    
    # For cut sheets, remove CS_ prefix
    if clean_name.upper().startswith('CS '):
        clean_name = clean_name[3:]
    
    # Consolidate multiple spaces into one and strip leading/trailing space
    clean_name = re.sub(r'\s+', ' ', clean_name).strip()
    
    return clean_name

def get_file_order_priority(file_type: str) -> int:
    """Get priority order for file types within a tag"""
    priority_order = {
        'Technical Data Sheet': 1,
        'Fan Curve': 2,
        'Drawing': 3,
        'Other': 4,
        'Item Summary': 5  # Should be filtered out anyway
    }
    return priority_order.get(file_type, 4)

def merge_user_customizations(fresh_data: Dict, existing_data: Dict) -> Dict:
    """
    Merge user customizations from existing structure into freshly generated structure.
    
    Preserves:
    - Custom display titles
    - Include/exclude flags
    - Custom positioning (if items still exist)
    
    Args:
        fresh_data: Newly generated structure data
        existing_data: Previous structure with user customizations
        
    Returns:
        Merged structure with user customizations applied
    """
    
    # Create lookup maps for existing items
    existing_structure = existing_data.get('pdf_structure', [])
    existing_lookup = {}
    
    # Build lookup by item identifier (type + tag + filename combination)
    for item in existing_structure:
        # Create unique identifier for each item
        identifier = f"{item.get('type', '')}|{item.get('tag', '')}|{item.get('filename', '')}"
        existing_lookup[identifier] = item
    
    # Apply customizations to fresh structure
    fresh_structure = fresh_data.get('pdf_structure', [])
    merged_structure = []
    
    for fresh_item in fresh_structure:
        # Find matching item in existing structure
        identifier = f"{fresh_item.get('type', '')}|{fresh_item.get('tag', '')}|{fresh_item.get('filename', '')}"
        existing_item = existing_lookup.get(identifier)
        
        if existing_item:
            # Preserve user customizations
            if 'display_title' in existing_item and existing_item['display_title'] != fresh_item.get('display_title'):
                fresh_item['display_title'] = existing_item['display_title']
                print(f"  Preserved custom title: {existing_item['display_title']}")
            
            if 'title' in existing_item and existing_item['title'] != fresh_item.get('title'):
                fresh_item['title'] = existing_item['title']
                print(f"  Preserved custom title: {existing_item['title']}")
            
            # Preserve include/exclude flags
            if 'include' in existing_item:
                fresh_item['include'] = existing_item['include']
                if not existing_item['include']:
                    print(f"  Preserved exclusion: {fresh_item.get('display_title', fresh_item.get('filename', 'Unknown'))}")
        
        merged_structure.append(fresh_item)
    
    # Update the fresh data with merged structure
    fresh_data['pdf_structure'] = merged_structure
    
    # Preserve metadata timestamp to indicate user edits exist
    if existing_data.get('metadata', {}).get('last_updated'):
        fresh_data['metadata']['last_updated'] = existing_data['metadata']['last_updated']
        fresh_data['metadata']['user_customized'] = True
    
    return fresh_data

def should_check_pricing_content(filename: str, use_filename_tags: bool) -> bool:
    """
    Determine if a file needs pricing content scanning for performance optimization.
    
    When using filename-based tag extraction, only "Item Summary" files contain
    pricing information that affects inclusion/exclusion decisions. All other document 
    types (Technical Data Sheet, Drawing, Fan Curve) don't need content scanning.
    
    Performance benefit: This optimization avoids unnecessary file I/O for 80%+ of files
    when using filename-based extraction, providing significant speed improvements.
    
    Args:
        filename: Name of file to check
        use_filename_tags: Whether filename-based extraction is enabled
        
    Returns:
        True if file needs pricing content check, False to skip (performance optimization)
    """
    if not use_filename_tags:
        # Original behavior: check all files when using content-based tag extraction
        # This is necessary because any file might contain pricing info when we don't
        # know the document type from filename patterns
        return True
    
    # When using filename tags, we can reliably identify document types from filenames
    # Only Item Summary files typically contain pricing information ($ symbols)
    # Technical Data Sheets, Drawings, Fan Curves rarely have pricing content
    filename_lower = filename.lower()
    return 'item summary' in filename_lower or 'item_summary' in filename_lower

def enhance_tag_mapping(tag_mapping: Dict[str, str], docs_path: str, no_pricing_filter: bool = False, use_filename_tags: bool = False, existing_user_edits: Dict = None) -> Dict:
    """
    Create complete PDF structure from tag mapping with optimized pricing filter.
    
    Args:
        tag_mapping: Dictionary of filename -> tag mappings
        docs_path: Path to documents directory  
        no_pricing_filter: If True, include all files regardless of pricing content
        use_filename_tags: If True, only scan Item Summary files for pricing (optimization)
        existing_user_edits: Optional existing structure with user customizations to preserve
        
    Returns:
        Enhanced mapping with complete PDF structure
    """
    
    # Log the start of structure generation
    log_processing_stage('enhance_tag_mapping', 'started', {
        'docs_path': docs_path,
        'no_pricing_filter': no_pricing_filter,
        'use_filename_tags': use_filename_tags,
        'tag_mapping_size': len(tag_mapping)
    })
    
    print("GENERATING COMPLETE PDF STRUCTURE")
    print("="*50)
    
    # Create the new structure
    pdf_structure = []
    position = 1
    
    # Group files by tag
    tag_groups = {}
    for filename, tag in tag_mapping.items():
        if tag:  # Only include files with tags
            if tag not in tag_groups:
                tag_groups[tag] = []
            tag_groups[tag].append(filename)
    
    # Sort tags (MAU first, then AHU, numerically)
    sorted_tags = sorted(tag_groups.keys(), key=lambda x: (
        x.startswith('AHU-'),  # MAU tags first (False comes before True)
        x.replace('AHU-', '').replace('MAU-', '').zfill(10)  # Then numerical order
    ))
    
    print(f"Processing {len(sorted_tags)} equipment tags: {sorted_tags}")
    
    # Log tag processing details
    log_processing_stage('tag_processing', 'started', {
        'sorted_tags': sorted_tags,
        'tag_count': len(sorted_tags),
        'tag_groups_size': len(tag_groups)
    })
    
    # Process each tag
    for tag in sorted_tags:
        files_for_tag = tag_groups[tag]
        
        # Check files for pricing content with smart optimization
        # When using filename tags, we only need to scan Item Summary files
        files_with_pricing_info = []
        scanned_count = 0
        skipped_count = 0
        
        for filename in files_for_tag:
            if should_check_pricing_content(filename, use_filename_tags):
                # This file needs pricing content scanning
                file_path = os.path.join(docs_path, filename)
                if has_pricing_content(file_path):
                    files_with_pricing_info.append(filename)
                    print(f"  [PRICING] Found pricing content in: {filename}")
                scanned_count += 1
            else:
                # Skip pricing check for performance (filename tags optimization)
                # Non-Item Summary files rarely contain pricing when using structured filenames
                skipped_count += 1
                
        # Log performance optimization results
        if use_filename_tags and skipped_count > 0:
            print(f"  [OPTIMIZATION] Scanned {scanned_count} files, skipped {skipped_count} (filename-based optimization)")
            
        # Log pricing analysis for this tag
        log_processing_stage('pricing_analysis', 'completed', {
            'tag': tag,
            'files_for_tag': len(files_for_tag),
            'files_with_pricing': len(files_with_pricing_info),
            'scanned_count': scanned_count,
            'skipped_count': skipped_count,
            'optimization_rate': (skipped_count / len(files_for_tag) * 100) if files_for_tag else 0
        })
        
        # Use all files for the tag (we'll mark pricing files as excluded if filter is enabled)
        filtered_files = files_for_tag
        
        # Add title page for this tag
        pdf_structure.append({
            "type": "title_page",
            "tag": tag,
            "title": tag,
            "position": position,
            "include": True
        })
        position += 1
        print(f"  [TITLE] Added title page for {tag} at position {position-1}")
        
        # Classify and sort files within the tag
        file_info = []
        for filename in filtered_files:
            file_type = classify_file_type(filename)
            display_title = create_display_title(filename)
            priority = get_file_order_priority(file_type)
            
            file_info.append({
                'filename': filename,
                'file_type': file_type,
                'display_title': display_title,
                'priority': priority
            })
        
        # Sort files by priority, then alphabetically
        file_info.sort(key=lambda x: (x['priority'], x['filename']))
        
        # Add documents for this tag
        for file_data in file_info:
            filename = file_data['filename']
            
            # Create converted path (will be set during actual conversion)
            base_name = os.path.splitext(filename)[0]
            converted_path = f"converted_pdfs/{base_name}.pdf"
            
            # Determine if this file should be included (not pricing file or pricing filter disabled)
            is_pricing_file = filename in files_with_pricing_info
            include_file = no_pricing_filter or not is_pricing_file
            
            pdf_structure.append({
                "type": "document",
                "tag": tag,
                "filename": filename,
                "display_title": file_data['display_title'],
                "file_type": file_data['file_type'],
                "converted_path": converted_path,
                "position": position,
                "include": include_file,
                "pricing_file": is_pricing_file
            })
            position += 1
            
            status = "[INCLUDED]" if include_file else "[EXCLUDED - PRICING]"
            print(f"    {status} {filename} -> {file_data['display_title']} (Type: {file_data['file_type']})")
    
    # Add cut sheets section
    cs_files = glob.glob(os.path.join(docs_path, "CS*.pdf"))
    if cs_files:
        # Add cut sheets title page
        pdf_structure.append({
            "type": "title_page",
            "tag": "CUT SHEETS",
            "title": "CUT SHEETS", 
            "position": position,
            "include": True
        })
        position += 1
        print(f"  [TITLE] Added CUT SHEETS title page at position {position-1}")
        
        # Add cut sheet files
        for cs_file in sorted(cs_files):
            filename = os.path.basename(cs_file)
            display_title = create_display_title(filename)
            
            pdf_structure.append({
                "type": "cut_sheet",
                "tag": "CUT SHEETS",  # Associate with CUT SHEETS title page
                "filename": filename,
                "display_title": display_title,
                "converted_path": cs_file,  # Cut sheets are already PDF
                "position": position,
                "include": True
            })
            position += 1
            print(f"    [CUT] {filename} -> {display_title}")
    
    # Create the complete data structure
    enhanced_data = {
        "pdf_structure": pdf_structure,
        "metadata": {
            "total_tags": len([item for item in pdf_structure if item["type"] == "title_page"]),
            "total_documents": len([item for item in pdf_structure if item["type"] == "document"]),
            "total_cut_sheets": len([item for item in pdf_structure if item["type"] == "cut_sheet"]),
            "total_items": len(pdf_structure),
            "processing_complete": True,
            "pricing_filter_enabled": not no_pricing_filter
        },
        # Keep legacy format for backward compatibility
        "tag_mapping": tag_mapping,
        "tag_groups": tag_groups
    }
    
    print(f"\nGenerated PDF structure with {len(pdf_structure)} total items:")
    print(f"  - {enhanced_data['metadata']['total_tags']} title pages")
    print(f"  - {enhanced_data['metadata']['total_documents']} documents") 
    print(f"  - {enhanced_data['metadata']['total_cut_sheets']} cut sheets")
    
    # Log completion of structure generation
    log_processing_stage('enhance_tag_mapping', 'completed', {
        'total_items': len(pdf_structure),
        'total_tags': enhanced_data['metadata']['total_tags'],
        'total_documents': enhanced_data['metadata']['total_documents'],
        'total_cut_sheets': enhanced_data['metadata']['total_cut_sheets'],
        'pricing_filter_enabled': enhanced_data['metadata']['pricing_filter_enabled']
    })
    
    # Merge existing user customizations if they exist
    if existing_user_edits:
        enhanced_data = merge_user_customizations(enhanced_data, existing_user_edits)
        print("Applied user customizations to generated structure")
    
    # Log JSON snapshot of enhanced structure
    log_json_snapshot('enhanced_structure', enhanced_data)
    
    return enhanced_data

def print_enhanced_summary(data: Dict):
    """Print enhanced summary of the PDF structure"""
    pdf_structure = data.get('pdf_structure', [])
    metadata = data.get('metadata', {})
    
    print("\n" + "="*60)
    print("COMPLETE PDF STRUCTURE SUMMARY")
    print("="*60)
    
    print(f"Total items in PDF: {metadata.get('total_items', len(pdf_structure))}")
    print(f"Title pages: {metadata.get('total_tags', 0)}")
    print(f"Documents: {metadata.get('total_documents', 0)}")
    print(f"Cut sheets: {metadata.get('total_cut_sheets', 0)}")
    print(f"Pricing filter enabled: {metadata.get('pricing_filter_enabled', True)}")
    
    print("\nComplete PDF Structure (in order):")
    current_tag = None
    for i, item in enumerate(pdf_structure, 1):
        if item["type"] == "title_page":
            current_tag = item["tag"]
            print(f"\n  {i:2d}. [TITLE PAGE] {item['title']}")
        elif item["type"] == "document":
            print(f"  {i:2d}.   +-- {item['display_title']} ({item['file_type']})")
        elif item["type"] == "cut_sheet":
            if current_tag != "CUT SHEETS":
                current_tag = "CUT SHEETS"
            print(f"  {i:2d}.   +-- {item['display_title']}")
    
    # Show any excluded items
    excluded_items = [item for item in pdf_structure if not item.get("include", True)]
    if excluded_items:
        print(f"\nExcluded items: {len(excluded_items)}")
        for item in excluded_items:
            print(f"  - {item.get('display_title', item.get('filename', 'Unknown'))}")

# Legacy function for backward compatibility
def create_tag_groups(tag_mapping: Dict[str, str]) -> Dict[str, list]:
    """Create tag groups from mapping (legacy compatibility)"""
    tag_groups = {}
    for filename, tag in tag_mapping.items():
        if tag:
            if tag not in tag_groups:
                tag_groups[tag] = []
            tag_groups[tag].append(filename)
    return tag_groups

if __name__ == "__main__":
    # Test the enhanced document processor
    print("Testing Enhanced Document Processor")
    print("="*50)
    
    # Example test data
    test_tag_mapping = {
        "10_Technical Data Sheet.docx": "AHU-1",
        "10_Fan Curve - Supply.jpg": "AHU-1", 
        "10_Drawing.doc": "AHU-1",
        "12_Technical Data Sheet.docx": "MAU-12",
        "12_Fan Curve - Supply.jpg": "MAU-12",
        "12_Item Summary.docx": "MAU-12",  # Should be filtered out
        "CS_Some_Product.pdf": None  # Cut sheet, no tag needed
    }
    
    test_docs_path = r"C:\Users\jacob\Claude\python-docx\documents\CS_Air_Handler_Light_Kit"
    
    # Test with pricing filter enabled
    enhanced_data = enhance_tag_mapping(test_tag_mapping, test_docs_path, no_pricing_filter=False)
    print_enhanced_summary(enhanced_data)
    
    # Save test results
    with open('test_pdf_structure.json', 'w', encoding='utf-8') as f:
        json.dump(enhanced_data, f, indent=2, ensure_ascii=False)
    
    print(f"\nTest results saved to: test_pdf_structure.json")